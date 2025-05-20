import streamlit as st
from typing import List, Dict, Any
import docx
from docx.shared import Inches
import plotly.graph_objects as go
import plotly.express as px
import pandas as pd
from datetime import datetime
import io
import base64
import re

class ReportBuilder:
    """Class for building and managing interactive reports"""
    
    def __init__(self):
        """Initialize the report builder"""
        if 'report_content' not in st.session_state:
            st.session_state.report_content = []
        if 'report_visualizations' not in st.session_state:
            st.session_state.report_visualizations = []
    
    def _parse_markdown_sections(self, content: str) -> List[Dict[str, Any]]:
        """Parse markdown content into sections based on headers
        
        Args:
            content (str): Markdown content with headers
            
        Returns:
            List[Dict]: List of sections with their content and level
        """
        sections = []
        current_section = {"title": "Introduction", "content": "", "level": 0}
        lines = content.split('\n')
        
        for line in lines:
            # Check for headers
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                # Save previous section if it has content
                if current_section["content"].strip():
                    sections.append(current_section)
                
                # Start new section
                level = len(header_match.group(1))
                title = header_match.group(2)
                current_section = {
                    "title": title,
                    "content": "",
                    "level": level
                }
            else:
                # Add line to current section
                current_section["content"] += line + "\n"
        
        # Add the last section
        if current_section["content"].strip():
            sections.append(current_section)
        
        return sections
    
    def add_content(self, content: str, title: str = None, level: int = 1):
        """Add content to the report
        
        Args:
            content (str): The content to add
            title (str, optional): Section title. Defaults to None.
            level (int, optional): Header level. Defaults to 1.
        """
        # Parse the content into sections
        sections = self._parse_markdown_sections(content)
        
        # If no sections were found, add the content as a single section
        if not sections:
            if title:
                st.session_state.report_content.append({
                    "type": "text",
                    "title": title,
                    "content": content,
                    "level": level
                })
            else:
                st.session_state.report_content.append({
                    "type": "text",
                    "title": "Section",
                    "content": content,
                    "level": level
                })
        else:
            # Add each section to the report
            for section in sections:
                st.session_state.report_content.append({
                    "type": "text",
                    "title": section["title"],
                    "content": section["content"].strip(),
                    "level": section["level"]
                })
    
    def add_visualization(self, fig: go.Figure, title: str = None):
        """Add a visualization to the report
        
        Args:
            fig (go.Figure): The Plotly figure to add
            title (str, optional): Visualization title. Defaults to None.
        """
        st.session_state.report_content.append({
            "type": "visualization",
            "title": title or "Visualization",
            "content": fig
        })
    
    def move_section(self, from_idx: int, to_idx: int):
        """Move a section in the report
        
        Args:
            from_idx (int): Index of section to move
            to_idx (int): Index to move section to
        """
        if 0 <= from_idx < len(st.session_state.report_content) and 0 <= to_idx < len(st.session_state.report_content):
            item = st.session_state.report_content.pop(from_idx)
            st.session_state.report_content.insert(to_idx, item)
    
    def remove_content(self, idx: int):
        """Remove content from the report
        
        Args:
            idx (int): Index of content to remove
        """
        if 0 <= idx < len(st.session_state.report_content):
            st.session_state.report_content.pop(idx)
    
    def export_to_word(self, title: str = "Report") -> bytes:
        """Export the report to a Word document
        
        Args:
            title (str): Title of the report
            
        Returns:
            bytes: The Word document as bytes
        """
        doc = docx.Document()
        
        # Add title
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # Add content
        for item in st.session_state.report_content:
            if item["type"] == "text":
                # Add section title with appropriate heading level
                doc.add_heading(item["title"], level=min(item["level"], 9))
                
                # Process the content to handle markdown formatting
                content = item["content"]
                # Split content into paragraphs
                paragraphs = content.split('\n\n')
                
                for para in paragraphs:
                    if para.strip():
                        # Create a new paragraph
                        p = doc.add_paragraph()
                        
                        # Split paragraph into runs (text between markdown formatting)
                        parts = re.split(r'(\*\*.*?\*\*)', para)
                        
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                # This is bold text
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            else:
                                # Regular text
                                p.add_run(part)
                
            elif item["type"] == "visualization":
                # Save the figure as an image
                fig = item["content"]
                # Update figure colors and layout before saving
                fig.update_layout(
                    template='plotly_white',
                    plot_bgcolor='white',
                    paper_bgcolor='white',
                    font=dict(color='black', size=12),
                    legend=dict(
                        orientation="h",
                        yanchor="bottom",
                        y=1.02,
                        xanchor="right",
                        x=1
                    ),
                    width=1200,  # Set larger width for better resolution
                    height=800,  # Set larger height for better resolution
                    margin=dict(t=100)  # Add top margin for legend
                )
                
                # Ensure colors are preserved
                for i, trace in enumerate(fig.data):
                    if hasattr(trace, 'line'):
                        trace.line.color = px.colors.qualitative.Set3[i % len(px.colors.qualitative.Set3)]
                    if hasattr(trace, 'marker'):
                        trace.marker.color = px.colors.qualitative.Set3[i % len(px.colors.qualitative.Set3)]
                
                # Export with higher resolution
                img_bytes = fig.to_image(
                    format="png",
                    width=1200,
                    height=800,
                    scale=2  # Increase scale for better resolution
                )
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def display_builder(self):
        """Display the report builder interface"""
        st.title("Report Builder")
        
        # Display current content
        if st.session_state.report_content:
            st.subheader("Current Report Content")
            
            for i, item in enumerate(st.session_state.report_content):
                st.markdown("---")
                
                if item["type"] == "text":
                    # Display section title
                    st.markdown(f"### {item['title']}")
                    
                    # Display content
                    st.markdown(item["content"])
                    
                    # Section controls
                    col1, col2, col3 = st.columns([1, 1, 1])
                    with col1:
                        if st.button("Move Up", key=f"move_up_{i}"):
                            if i > 0:
                                self.move_section(i, i-1)
                                st.rerun()
                    with col2:
                        if st.button("Move Down", key=f"move_down_{i}"):
                            if i < len(st.session_state.report_content) - 1:
                                self.move_section(i, i+1)
                                st.rerun()
                    with col3:
                        if st.button("Remove Section", key=f"remove_{i}"):
                            self.remove_content(i)
                            st.rerun()
                
                elif item["type"] == "visualization":
                    # Display visualization
                    st.plotly_chart(item["content"], use_container_width=True, key=f"report_chart_{i}")
                    
                    # Section controls
                    col1, col2, col3 = st.columns([1, 1, 1])
                    with col1:
                        if st.button("Move Up", key=f"move_up_viz_{i}"):
                            if i > 0:
                                self.move_section(i, i-1)
                                st.rerun()
                    with col2:
                        if st.button("Move Down", key=f"move_down_viz_{i}"):
                            if i < len(st.session_state.report_content) - 1:
                                self.move_section(i, i+1)
                                st.rerun()
                    with col3:
                        if st.button("Remove Visualization", key=f"remove_viz_{i}"):
                            self.remove_content(i)
                            st.rerun()
        
        # Export options
        st.subheader("Export Report")
        export_format = st.selectbox(
            "Select export format",
            options=["Word Document", "PDF"],
            key="export_format_select"
        )
        
        if st.button("Export Report", key="export_button"):
            if export_format == "Word Document":
                doc_bytes = self.export_to_word()
                st.download_button(
                    label="Download Word Document",
                    data=doc_bytes,
                    file_name="report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_word"
                )
            else:  # PDF
                # TODO: Implement PDF export
                st.info("PDF export coming soon!") 